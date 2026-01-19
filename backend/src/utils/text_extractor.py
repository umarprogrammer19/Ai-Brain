import fitz  # PyMuPDF
from pypdf import PdfReader
from docx import Document
from io import BytesIO
import os


class TextExtractor:
    """
    Utility class to extract text from various document formats
    Supported formats: PDF, DOCX, TXT, MD
    """

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """
        Extract text from PDF file using PyMuPDF (fitz) for better reliability

        Args:
            file_bytes: Raw bytes of the PDF file

        Returns:
            str: Extracted text from the PDF
        """
        text = ""
        try:
            # Using PyMuPDF (fitz) to open from bytes stream
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            doc.close()
        except Exception as e:
            print(f"Error extracting text from PDF with PyMuPDF: {e}")
            # Fallback to PyPDF2/pypdf
            try:
                pdf_stream = BytesIO(file_bytes)
                pdf_reader = PdfReader(pdf_stream)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            except Exception as e2:
                print(f"Error extracting text from PDF with pypdf: {e2}")

        return text

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """
        Extract text from DOCX file

        Args:
            file_bytes: Raw bytes of the DOCX file

        Returns:
            str: Extracted text from the DOCX
        """
        text = ""
        try:
            docx_stream = BytesIO(file_bytes)
            doc = Document(docx_stream)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")

        return text

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        """
        Extract text from TXT file

        Args:
            file_bytes: Raw bytes of the TXT file

        Returns:
            str: Extracted text from the TXT
        """
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Try different encodings
            try:
                return file_bytes.decode('latin-1')
            except UnicodeDecodeError:
                return file_bytes.decode('utf-8', errors='ignore')

    @staticmethod
    def extract_text_from_md(file_bytes: bytes) -> str:
        """
        Extract text from MD (Markdown) file

        Args:
            file_bytes: Raw bytes of the MD file

        Returns:
            str: Extracted text from the MD
        """
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_bytes.decode('latin-1')
            except UnicodeDecodeError:
                return file_bytes.decode('utf-8', errors='ignore')

    @classmethod
    def extract_text(cls, file_bytes: bytes, file_extension: str) -> str:
        """
        Extract text based on file extension

        Args:
            file_bytes: Raw bytes of the file
            file_extension: Extension of the file ('.pdf', '.docx', '.txt', '.md')

        Returns:
            str: Extracted text from the file
        """
        extractor_map = {
            '.pdf': cls.extract_text_from_pdf,
            '.docx': cls.extract_text_from_docx,
            '.txt': cls.extract_text_from_txt,
            '.md': cls.extract_text_from_md
        }

        if file_extension.lower() not in extractor_map:
            raise ValueError(f"Unsupported file type: {file_extension}")

        extractor_func = extractor_map[file_extension.lower()]
        return extractor_func(file_bytes)